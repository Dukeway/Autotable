import os
import pandas as pd
import logging
from docx import Document
import json
from datetime import datetime

logger = logging.getLogger(__name__)

class AutoTable:
    """自动化填表处理核心类"""
    def __init__(self, knowledge_base_path, word_template_path, llm_client, output_folder="output"):
        self.knowledge_base_path = knowledge_base_path
        self.word_template_path = word_template_path
        self.output_folder = output_folder
        self.llm_client = llm_client
        self.knowledge_base = None
        self.knowledge_dict = None
        self.doc = None

        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

    def load_knowledge_base(self):
        try:
            logger.info(f"正在加载知识库: {self.knowledge_base_path}")
            self.knowledge_base = pd.read_excel(self.knowledge_base_path)
            required_columns = ['字段', '字段值']
            if not all(col in self.knowledge_base.columns for col in required_columns):
                missing = [col for col in required_columns if col not in self.knowledge_base.columns]
                logger.error(f"知识库缺少必要列: {missing}")
                return False
            self.knowledge_dict = dict(zip(self.knowledge_base['字段'], self.knowledge_base['字段值']))
            logger.info(f"知识库加载完成，共{len(self.knowledge_dict)}个字段")
            return True
        except Exception as e:
            logger.error(f"知识库加载异常: {str(e)}")
            return False

    def load_template(self):
        try:
            logger.info(f"正在加载Word模板: {self.word_template_path}")
            self.doc = Document(self.word_template_path)
            logger.info(f"模板加载完成，包含{len(self.doc.tables)}个表格")
            return True
        except Exception as e:
            logger.error(f"模板加载失败: {str(e)}")
            return False

    def analyze_tables_with_llm(self, table):
        table_text = []
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                table_text.append(f"({row_idx},{cell_idx}): {cell.text.strip()}")
        prompt = f"""
        请分析以下表格结构，识别字段名称及其对应填写位置：
        {table_text}

        返回JSON格式结果，包含字段名及其对应的单元格位置：
        {{
            "字段名1": {{"field_cell": [行,列], "value_cell": [行,列]}},
            "字段名2": {{"field_cell": [行,列], "value_cell": [行,列]}}
        }}
        请确保只返回JSON格式数据，不要包含其他内容。
        """
        try:
            messages = [
                {"role": "system", "content": "你是一个专业的文档结构分析助手"},
                {"role": "user", "content": prompt}
            ]
            result = self.llm_client.chat_completion(messages, temperature=0.1)
            json_str = self._extract_json(result)
            return json.loads(json_str)
        except Exception as e:
            logger.error(f"表格分析失败: {str(e)}")
            return {}

    def _extract_json(self, text):
        try:
            json.loads(text)
            return text
        except json.JSONDecodeError:
            import re
            json_match = re.search(r'({.*})', text.replace('\n', ''), re.DOTALL)
            if json_match:
                return json_match.group(1)
            raise ValueError("未找到有效JSON内容")

    def fill_tables(self):
        if not self.doc or not self.knowledge_dict:
            logger.error("文档或知识库未正确初始化")
            return False
        filled_count = 0
        missing_fields = []
        for table_idx, table in enumerate(self.doc.tables):
            logger.info(f"正在处理第{table_idx + 1}个表格")
            fields_map = self.analyze_tables_with_llm(table)
            for field, positions in fields_map.items():
                if field in self.knowledge_dict:
                    try:
                        row, col = positions["value_cell"]
                        table.cell(row, col).text = str(self.knowledge_dict[field])
                        filled_count += 1
                        logger.debug(f"已填写字段: {field}")
                    except IndexError:
                        logger.error(f"无效的单元格位置: {positions}")
                    except Exception as e:
                        logger.error(f"字段填写异常: {field} - {str(e)}")
                else:
                    missing_fields.append(field)
                    logger.warning(f"未找到字段定义: {field}")
        logger.info(f"完成表格填充，共填写{filled_count}个字段")
        if missing_fields:
            logger.warning(f"缺失字段列表: {', '.join(missing_fields)}")
        return True

    def save_document(self, filename=None):
        if not self.doc:
            logger.error("文档实例未初始化")
            return False
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = os.path.splitext(os.path.basename(self.word_template_path))[0]
            filename = f"{base_name}_filled_{timestamp}.docx"
        output_path = os.path.join(self.output_folder, filename)
        try:
            self.doc.save(output_path)
            logger.info(f"文档已保存至: {output_path}")
            return True
        except Exception as e:
            logger.error(f"文档保存失败: {str(e)}")
            return False

    def run(self):
        logger.info("启动自动化填表流程")
        if all([
            self.load_knowledge_base(),
            self.load_template(),
            self.fill_tables(),
            self.save_document()
        ]):
            logger.info("流程执行成功")
            return True
        logger.error("流程执行过程中发生错误")
        return False